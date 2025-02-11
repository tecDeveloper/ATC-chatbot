from docx import Document

def add_link_to_answers(input_file, output_file, link):
    doc = Document(input_file)
    new_doc = Document()
    
    paragraphs = doc.paragraphs  # Store paragraphs in a list
    total_paragraphs = len(paragraphs)

    for i, para in enumerate(paragraphs):
        words = para.text.split()
        new_doc.add_paragraph(para.text)

        # Check if the next paragraph starts with "QUESTION"
        if any(word == "QUESTION:" for word in words) and i > 0:
            new_doc.add_paragraph(f"You can further read about it here: {link}")

    new_doc.save(output_file)
    print(f"Updated document saved as {output_file}")


# Usage
input_docx = "./Chatbot Preprocessed Questions.docx"  
output_docx = "Updated.docx"
link = "https://pub-2bae66d5e6d74bfda26d9e7d8ee03534.r2.dev/documents/ATCMarket%20Free%20Membership%20Agreement.pdf"

add_link_to_answers(input_docx, output_docx, link)
